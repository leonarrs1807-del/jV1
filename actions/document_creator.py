import os
from pathlib import Path
from datetime import datetime

def document_creator(parameters: dict, player=None) -> str:
    """
    Crea documentos de texto, Word o Excel en base a los parámetros.
    """
    action = parameters.get("action", "").lower()
    title = parameters.get("title", "Documento_Sin_Titulo")
    content = parameters.get("content", "")
    sheets = parameters.get("sheets", [])
    
    # Determinar ruta de guardado (por defecto, el Escritorio)
    desktop_path = Path(os.path.join(os.environ["USERPROFILE"], "Desktop"))
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = "".join([c for c in title if c.isalpha() or c.isdigit() or c==' ']).rstrip().replace(" ", "_")
    if not safe_title:
        safe_title = "Documento"
        
    try:
        if action == "word" or action == "google_doc":
            # Si piden google doc, por ahora lo hacemos Word local y avisamos.
            try:
                from docx import Document
                doc = Document()
                doc.add_heading(title, 0)
                
                # Procesamiento simple de markdown a Word
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    if line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('# '):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith('- '):
                        doc.add_paragraph(line[2:], style='List Bullet')
                    else:
                        doc.add_paragraph(line)
                        
                file_name = f"{safe_title}_{timestamp}.docx"
                file_path = desktop_path / file_name
                doc.save(file_path)
                return f"Documento Word creado exitosamente en tu Escritorio como '{file_name}'."
            except ImportError:
                return "Error: Faltan librerías para crear Word. (python-docx)"
                
        elif action == "excel" or action == "google_sheet":
            try:
                from openpyxl import Workbook
                wb = Workbook()
                wb.remove(wb.active) # Remove default sheet
                
                if not sheets:
                    return "Error: No se proporcionaron datos (sheets) para crear el Excel."
                    
                for sheet_data in sheets:
                    sheet_name = sheet_data.get("name", "Hoja")
                    headers = sheet_data.get("headers", [])
                    rows = sheet_data.get("rows", [])
                    
                    ws = wb.create_sheet(title=sheet_name[:31])
                    if headers:
                        ws.append(headers)
                    for row in rows:
                        ws.append(row)
                        
                file_name = f"{safe_title}_{timestamp}.xlsx"
                file_path = desktop_path / file_name
                wb.save(file_path)
                return f"Planilla Excel creada exitosamente en tu Escritorio como '{file_name}'."
            except ImportError:
                return "Error: Faltan librerías para crear Excel. (openpyxl)"
                
        elif action == "text":
            file_name = f"{safe_title}_{timestamp}.txt"
            file_path = desktop_path / file_name
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(f"{title}\n\n{content}")
            return f"Archivo de texto creado en tu Escritorio como '{file_name}'."
            
        else:
            return f"Acción '{action}' no soportada o desconocida. Usá 'word', 'excel' o 'text'."
            
    except Exception as e:
        return f"Error al crear el documento: {str(e)}"
